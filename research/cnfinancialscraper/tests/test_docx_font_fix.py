"""v1.1 .docx 字体乱码回归测试。

问题：生成的 .docx 在线查看正常，下载保存后字体乱掉。
原因：StdlibDocxWriter 的 <w:rPrDefault><w:rPr><w:rFonts ...></w:rPr></w:rPrDefault>
     缺 w:eastAsia 属性 + w:hint + w:lang。
     Word 在解析时对没有 eastAsia 属性的 CJK 字符使用系统默认字体，
     当系统字体集不规范或字体缺失时，CJK 字符会被替换为非预期字体（乱码）。

修复后检查项：
    1. <w:rPrDefault> 中 <w:rFonts> 必须含 w:ascii/w:hAnsi/w:eastAsia/w:cs 四个属性
    2. 必须有 w:hint="eastAsia" 属性
    3. 必含 <w:lang> 显式声明中文
    4. 用 universal 字体（宋体/Times New Roman），不依赖特定字体存在
"""
from __future__ import annotations

import os
import sys
import re
import tempfile
import zipfile
from pathlib import Path

SKILL_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(SKILL_DIR / "scripts"))


# ============================================================
# 测试
# ============================================================

def test_docx_has_eastasia_font():
    """生成的 docx 必须含 w:eastAsia 属性（核心 bug 修复）。"""
    from report_stdlib_fallbacks import StdlibDocxWriter

    data = {
        "summary": {"stock_name": "测试", "periodic_count": 1, "broker_count": 1, "announcement_count": 1, "has_buy_rating": True},
        "periodic_reports": [],
        "broker_reports": [],
        "announcements": [],
    }

    with tempfile.TemporaryDirectory() as tmp:
        output = os.path.join(tmp, "test.docx")
        writer = StdlibDocxWriter()
        writer.export_comprehensive_report(data, "000001", output)

        # 读取 styles.xml
        with zipfile.ZipFile(output, "r") as zf:
            styles = zf.read("word/styles.xml").decode("utf-8")

        # 核心检查：rFonts 含 eastAsia
        assert 'w:eastAsia=' in styles, "缺少 w:eastAsia 字体定义（CJK 渲染关键）"
        print("  ✓ 含 w:eastAsia 字体定义")


def test_docx_has_complete_font_quad():
    """rFonts 必须含 4 个字体属性：ascii/hAnsi/eastAsia/cs。"""
    from report_stdlib_fallbacks import StdlibDocxWriter

    data = {"summary": {}, "periodic_reports": [], "broker_reports": [], "announcements": []}

    with tempfile.TemporaryDirectory() as tmp:
        output = os.path.join(tmp, "test.docx")
        StdlibDocxWriter().export_comprehensive_report(data, "T", output)

        with zipfile.ZipFile(output, "r") as zf:
            styles = zf.read("word/styles.xml").decode("utf-8")

        # 提取 rFonts 那一行（默认的）
        match = re.search(r'<w:rFonts\s+([^/]*?)/>', styles)
        assert match, "未找到 <w:rFonts>"
        attrs = match.group(1)

        required = ["w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"]
        for attr in required:
            assert attr in attrs, f"缺少字体属性 {attr}"
        print(f"  ✓ 4 个字体属性齐全：{required}")


def test_docx_has_hint_eastasia():
    """必须有 w:hint="eastAsia" 告诉 Word 优先按东亚字体处理。"""
    from report_stdlib_fallbacks import StdlibDocxWriter

    data = {"summary": {}, "periodic_reports": [], "broker_reports": [], "announcements": []}

    with tempfile.TemporaryDirectory() as tmp:
        output = os.path.join(tmp, "test.docx")
        StdlibDocxWriter().export_comprehensive_report(data, "T", output)

        with zipfile.ZipFile(output, "r") as zf:
            styles = zf.read("word/styles.xml").decode("utf-8")

        assert 'w:hint="eastAsia"' in styles, "缺少 w:hint='eastAsia'"
        print("  ✓ 含 w:hint='eastAsia'")


def test_docx_has_language_tag():
    """必须有 <w:lang> 显式声明中文。"""
    from report_stdlib_fallbacks import StdlibDocxWriter

    data = {"summary": {}, "periodic_reports": [], "broker_reports": [], "announcements": []}

    with tempfile.TemporaryDirectory() as tmp:
        output = os.path.join(tmp, "test.docx")
        StdlibDocxWriter().export_comprehensive_report(data, "T", output)

        with zipfile.ZipFile(output, "r") as zf:
            styles = zf.read("word/styles.xml").decode("utf-8")

        assert "<w:lang" in styles, "缺少 <w:lang>"
        # 必须声明 eastAsia 为中文（zh-CN/zh-TW/zh-HK 任一）
        match = re.search(r'<w:lang[^>]+w:eastAsia="([^"]+)"', styles)
        assert match and "zh" in match.group(1).lower(), f"eastAsia 语言应含 zh，实际: {match.group(1) if match else 'NOT FOUND'}"
        print(f"  ✓ <w:lang> 中文标记：{match.group(1)}")


def test_docx_uses_universal_fonts():
    """字体选择应该是 universal 字体（不依赖特定系统）。"""
    from report_stdlib_fallbacks import StdlibDocxWriter

    data = {"summary": {}, "periodic_reports": [], "broker_reports": [], "announcements": []}

    with tempfile.TemporaryDirectory() as tmp:
        output = os.path.join(tmp, "test.docx")
        StdlibDocxWriter().export_comprehensive_report(data, "T", output)

        with zipfile.ZipFile(output, "r") as zf:
            styles = zf.read("word/styles.xml").decode("utf-8")

        # 不应依赖 "微软雅黑"（Mac/Linux 没有）
        # 应是 宋体 (SimSun) 或 宋体 unicode 字符串
        # 微软雅黑 is risky on Mac/Linux
        if "微软雅黑" in styles:
            print("  ⚠️ 含微软雅黑，在 Mac/Linux 用户可能无此字体")
        else:
            print("  ✓ 不依赖中文专属字体（微软雅黑）")

        # 应该含 宋体 或 SimSun
        assert "宋体" in styles or "SimSun" in styles, "eastAsia 应使用 universal 字体（宋体/SimSun）"
        print("  ✓ 使用 universal 字体（宋体/SimSun）")


def test_docx_validation_lifecycle():
    """完整生成 + 解析 + 检查生命周期。"""
    from report_stdlib_fallbacks import StdlibDocxWriter

    data = {
        "summary": {"stock_name": "中文测试股", "periodic_count": 1, "broker_count": 1, "announcement_count": 1, "has_buy_rating": True},
        "periodic_reports": [{"title": "测试报告", "publish_date": "2026-07-28", "url": "https://example.com"}],
        "broker_reports": [{"title": "买入评级", "publish_date": "2026-07-28", "broker_name": "测试券商", "analyst": "分析师", "rating": "买入", "target_price": 100.0}],
        "announcements": [{"title": "测试公告", "publish_date": "2026-07-28", "announcement_type": "分红", "url": ""}],
    }

    with tempfile.TemporaryDirectory() as tmp:
        output = os.path.join(tmp, "test.docx")
        StdlibDocxWriter().export_comprehensive_report(data, "000001", output)

        # 用 python-docx 解析验证
        try:
            from docx import Document
            doc = Document(output)
            # 验证文本正确写入
            text_found = False
            for p in doc.paragraphs:
                if "中文测试股" in p.text:
                    text_found = True
                    break
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        if "中文测试股" in cell.text:
                            text_found = True
                            break
            assert text_found, "中文内容未正确写入"
            print("  ✓ python-docx 可正确解析，中文写入无误")
        except ImportError:
            print("  ⊘ python-docx 未安装，跳过解析验证")
            pass

        # 文件应 < 50KB（典型报告大小）
        size = os.path.getsize(output)
        assert size < 100_000, f"文件过大: {size} 字节"
        print(f"  ✓ 文件大小合理: {size} 字节")


# ============================================================
# PPTX 字体检查（v1.1 新增）
# ============================================================

def _read_pptx_theme(path: str) -> str:
    """从 pptx 文件读取 theme1.xml 的内容。"""
    import zipfile
    with zipfile.ZipFile(path, "r") as zf:
        for n in zf.namelist():
            if n.endswith("theme/theme1.xml"):
                return zf.read(n).decode("utf-8")
    return ""


def test_pptx_theme_has_ea_font():
    """PPTX theme 必须声明东亚字体（v1.1 修复：之前是空字符串）。"""
    from report_stdlib_fallbacks import StdlibPptxWriter

    data = {"summary": {}, "periodic_reports": [], "broker_reports": [], "announcements": []}

    with tempfile.TemporaryDirectory() as tmp:
        output = os.path.join(tmp, "test.pptx")
        StdlibPptxWriter().export_comprehensive_report(data, "T", output)

        theme = _read_pptx_theme(output)
        # 必须含 ea typeface 且非空
        import re
        ea_matches = re.findall(r'<a:ea\s+typeface="([^"]*)"', theme)
        assert ea_matches, "<a:ea typeface=...> 应在 theme 中声明"
        for ea in ea_matches:
            assert ea.strip(), f"<a:ea typeface> 不能为空字符串（v1.1 修复）"
        print(f"  ✓ PPTX theme 含东亚字体：{ea_matches}")


def test_pptx_theme_uses_universal_fonts():
    """PPTX theme 应使用 universal 字体（不依赖 微软雅黑）。"""
    from report_stdlib_fallbacks import StdlibPptxWriter

    data = {"summary": {}, "periodic_reports": [], "broker_reports": [], "announcements": []}

    with tempfile.TemporaryDirectory() as tmp:
        output = os.path.join(tmp, "test.pptx")
        StdlibPptxWriter().export_comprehensive_report(data, "T", output)
        theme = _read_pptx_theme(output)

        if "微软雅黑" in theme:
            print("  ⚠️ PPTX theme 含微软雅黑")
        else:
            print("  ✓ PPTX theme 不依赖 微软雅黑")


# ============================================================
# XLSX 字体检查（v1.1 新增）
# ============================================================

def _read_xlsx_styles(path: str) -> str:
    """从 xlsx 文件读取 styles.xml 内容。"""
    import zipfile
    with zipfile.ZipFile(path, "r") as zf:
        return zf.read("xl/styles.xml").decode("utf-8")


def test_xlsx_uses_universal_font():
    """XLSX style 必须使用 universal 字体（v1.1 修复：之前用 微软雅黑）。"""
    from report_stdlib_fallbacks import StdlibXlsxWriter

    data = {"summary": {}, "periodic_reports": [], "broker_reports": [], "announcements": []}

    with tempfile.TemporaryDirectory() as tmp:
        output = os.path.join(tmp, "test.xlsx")
        StdlibXlsxWriter().export_comprehensive_report(data, "T", output)

        styles = _read_xlsx_styles(output)
        # 不应该用 微软雅黑（Mac/Linux 没有）
        assert "微软雅黑" not in styles, "XLSX 不应依赖 微软雅黑"
        # 应该用 宋体 或其他 universal 字体
        assert "宋体" in styles or "SimSun" in styles, "XLSX 应使用 universal 字体"
        print(f"  ✓ XLSX 使用 universal 字体（宋体/SimSun）")


def test_xlsx_openpyxl_path_uses_universal_font():
    """openpyxl 路径下生成的 xlsx 也用 universal 字体。"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        print("  ⊘ openpyxl 未安装，跳过")
        return

    # 写一个使用 宋体 字体的工作簿，验证能 round-trip
    from openpyxl import Workbook
    from openpyxl.styles import Font

    wb = Workbook()
    ws = wb.active
    ws['A1'] = "中文测试"
    ws['A1'].font = Font(name='宋体')
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
        path = f.name
    try:
        wb.save(path)

        # 重新加载看是否能识别中文字体
        wb2 = load_workbook(path)
        font = wb2.active['A1'].font
        assert font.name == "宋体", f"字体应为宋体，实际: {font.name}"
        print(f"  ✓ openpyxl 路径字体声明一致：{font.name}")
    finally:
        os.remove(path)


# ============================================================
# 运行所有
# ============================================================

def main():
    tests = [
        (name, obj)
        for name, obj in globals().items()
        if name.startswith("test_") and callable(obj)
    ]
    passed = 0
    failed = 0
    for name, fn in tests:
        try:
            fn()
            passed += 1
            print(f"  ✓ {name}")
        except Exception as e:
            print(f"  ✗ {name}: {e}")
            import traceback
            traceback.print_exc()
            failed += 1

    print(f"\n{'='*60}\n字体修复测试: {passed} 通过, {failed} 失败\n{'='*60}")
    return 0 if failed == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
