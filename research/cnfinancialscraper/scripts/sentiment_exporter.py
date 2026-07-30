# -*- coding: utf-8 -*-
"""
全网舆情爬虫 — 多格式导出器 v4.3
================================
把 SentimentSnapshot 导出为：
  1. 对话提示（dialog）    — 直接贴入对话窗口
  2. Word (.docx)         — 含封面 / 目录 / 分类 / 摘要，正面 vs 舆情分章
  3. Excel (.xlsx)        — 分sheet 正面 / 舆情 / 中性 / 全部
  4. JSON / CSV           — 机器可读备份

输出位置：默认 data/sentiment_exports/
"""
from __future__ import annotations

import csv
import json
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime

# 兼容可选依赖
try:
    from openpyxl import Workbook  # type: ignore
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side  # type: ignore
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from docx import Document  # type: ignore
    from docx.shared import Pt, Cm, RGBColor  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from sentiment_crawler import SentimentSnapshot, SentimentArticle  # type: ignore
except ImportError:
    SentimentSnapshot = SentimentArticle = None  # type: ignore

# 路径
SKILL_DATA_DIR = Path(__file__).resolve().parent.parent / "data"
EXPORT_DIR = SKILL_DATA_DIR / "sentiment_exports"
EXPORT_DIR.mkdir(parents=True, exist_ok=True)

logger = logging.getLogger("sentiment_exporter")
logger.setLevel(logging.INFO)


# ================================================================
# 1. 对话提示 — 给大模型/用户直接看
# ================================================================

def to_dialog(snapshot: SentimentSnapshot,
              top_k: int = 30,
              with_links: bool = True) -> str:
    """生成可直接贴到对话的精简提示"""
    if snapshot is None or not snapshot.articles:
        return "📭 本次全网舆情爬取未获取到任何结果，请尝试放宽条件。"

    lines: List[str] = []
    lines.append(f"📰 **全网舆情快照** `{snapshot.snapshot_id}`")
    lines.append(f"📅 抓取时间：{snapshot.created_at}")
    lines.append(
        f"📊 共 {len(snapshot.articles)} 条 | 正面 {snapshot.positive_count()} | "
        f"舆情 {snapshot.negative_count()} | 中性 {snapshot.neutral_count()}"
    )
    sev = snapshot.stats.get("by_severity", {})
    if sev:
        sev_str = " | ".join(f"{k}: {v}" for k, v in sev.items() if k != "中性")
        if sev_str:
            lines.append(f"⚠️ 风险分布：{sev_str}")
    # 媒体来源 Top 5
    src_dist = snapshot.stats.get("by_source_type", {})
    if src_dist:
        src_top = sorted(src_dist.items(), key=lambda x: -x[1])[:5]
        src_str = " | ".join(f"{k}:{v}" for k, v in src_top)
        lines.append(f"🗞️ 媒体类别 Top: {src_str}")
    # 耗时
    elapsed = snapshot.stats.get("elapsed_seconds")
    if elapsed is not None:
        lines.append(f"⏱ 耗时 {elapsed:.1f}s" + (" (超时返回)" if snapshot.stats.get("timed_out") else ""))
    lines.append("")

    # 正面
    pos = [a for a in snapshot.articles if a.sentiment == "positive"][:top_k]
    if pos:
        lines.append(f"### ✅ 正面新闻 ({len(pos)})")
        for i, a in enumerate(pos, 1):
            lines.append(_dialog_one_line(i, a, with_links))
        lines.append("")

    # 舆情（按严重度）
    negs = [a for a in snapshot.articles if a.sentiment == "negative"]
    negs.sort(key=lambda a: (-a.sentiment_score, a.publish_time or ""))
    if negs:
        lines.append(f"### ⚠️ 舆情（负面）({len(negs)})")
        for i, a in enumerate(negs[:top_k], 1):
            lines.append(_dialog_one_line(i, a, with_links))
        lines.append("")

    # 中性
    neu = [a for a in snapshot.articles if a.sentiment == "neutral"][:max(5, top_k // 3)]
    if neu:
        lines.append(f"### ℹ️ 中性公告 ({len(neu)})")
        for i, a in enumerate(neu, 1):
            lines.append(_dialog_one_line(i, a, with_links))
        lines.append("")

    lines.append("---")
    lines.append(f"📁 完整快照保存于：{snapshot.extra_path if hasattr(snapshot, 'extra_path') else 'data/sentiment_snapshots/'}")  # type: ignore[attr-defined]
    lines.append(f"💡 提示：可对我说「导出Word」「导出Excel」「按严重度筛选」获得更多处理。")
    return "\n".join(lines)


def _dialog_one_line(idx: int, a: SentimentArticle, with_links: bool) -> str:
    risk_icon = {
        "低度关注": "🟡", "中度舆情": "🟠", "高危舆情": "🔴",
        "低度利好": "🟢", "中度利好": "🟢", "重大利好": "✅",
    }.get(a.severity, "⚪")
    pos_icon = {"positive": "🟢", "negative": "🔴", "neutral": "⚪"}.get(a.sentiment, "⚪")
    fields = [
        f"{idx}. {pos_icon}{risk_icon} **[{a.target_name}]**",
        f"《{a.title}》",
        f"  · 平台：{a.source or '未知'} ({a.source_type or '媒体'})",
        f"  · 时间：{a.publish_time or '未知'}",
        f"  · 摘要：{a.summary or '(无摘要)'}",
    ]
    if with_links and a.url:
        fields.append(f"  · 链接：{a.url}")
    if a.keywords_matched:
        kw_text = "、".join(a.keywords_matched[:6])
        fields.append(f"  · 命中关键词：{kw_text}")
    return "\n".join(fields)


# ================================================================
# 2. Excel (.xlsx) — 包含总览 + 分类
# ================================================================

def to_excel(snapshot: SentimentSnapshot,
             output_path: Optional[Path] = None) -> str:
    """导出为 xlsx 多 sheet。
    Sheet:
        概览 / 正面新闻 / 舆情（负面） / 中性 / 全部
    """
    if not HAS_OPENPYXL:
        return _placeholder("Excel 导出需要 openpyxl: pip install openpyxl")
    if snapshot is None:
        return _placeholder("快照为空")

    output_path = Path(output_path or EXPORT_DIR / f"{snapshot.snapshot_id}.xlsx")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    wb = Workbook()

    # ---- Sheet 1 概览 ----
    ws = wb.active
    ws.title = "概览"
    ws["A1"] = "全网舆情快照"
    ws["A1"].font = Font(size=18, bold=True)
    ws.merge_cells("A1:E1")

    rows = [
        ("快照ID", snapshot.snapshot_id),
        ("抓取时间", snapshot.created_at),
        ("目标过滤", json.dumps(snapshot.target_filter, ensure_ascii=False)),
        ("媒体过滤", " / ".join(snapshot.source_filter)),
        ("", ""),
        ("总条数", snapshot.stats.get("total", 0)),
        ("正面", snapshot.positive_count()),
        ("舆情", snapshot.negative_count()),
        ("中性", snapshot.neutral_count()),
        ("", ""),
        ("风险分布", json.dumps(snapshot.stats.get("by_severity", {}), ensure_ascii=False)),
        ("媒体类别分布", json.dumps(snapshot.stats.get("by_source_type", {}), ensure_ascii=False)),
    ]
    for i, (k, v) in enumerate(rows, 3):
        ws.cell(row=i, column=1, value=k).font = Font(bold=True)
        ws.cell(row=i, column=2, value=v)
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 80

    # ---- Sheet 2-5 分类明细 ----
    for sheet_name, sentiment in [
        ("正面新闻", "positive"),
        ("舆情负面", "negative"),
        ("中性公告", "neutral"),
        ("全部文章", "all"),
    ]:
        ws = wb.create_sheet(sheet_name)
        headers = ["序号", "标题", "内容简介", "发布平台", "媒体类型", "发布时间", "页面连接", "目标", "严重等级", "情感", "命中关键词"]
        for col, h in enumerate(headers, 1):
            c = ws.cell(row=1, column=col, value=h)
            c.font = Font(bold=True, color="FFFFFF", name="宋体")  # v4.3.2: 添加中文字体
            c.fill = PatternFill("solid", fgColor="2E5984")
            c.alignment = Alignment(horizontal="center")
        items = snapshot.articles if sentiment == "all" else [a for a in snapshot.articles if a.sentiment == sentiment]
        items.sort(key=lambda a: (-(1 if a.severity in ("高危舆情", "中度舆情") else 0), a.publish_time or ""))
        for r, a in enumerate(items, 2):
            ws.cell(row=r, column=1, value=r - 1)
            ws.cell(row=r, column=2, value=a.title)
            ws.cell(row=r, column=3, value=a.summary)
            ws.cell(row=r, column=4, value=a.source)
            ws.cell(row=r, column=5, value=a.source_type)
            ws.cell(row=r, column=6, value=a.publish_time)
            cell_url = ws.cell(row=r, column=7, value=a.url)
            cell_url.hyperlink = a.url or None
            cell_url.font = Font(color="0563C1", underline="single")
            ws.cell(row=r, column=8, value=f"{a.target_name} ({a.target_type})")
            sev_cell = ws.cell(row=r, column=9, value=a.severity)
            if a.severity == "高危舆情":
                sev_cell.fill = PatternFill("solid", fgColor="FFC7CE")
                sev_cell.font = Font(bold=True, color="9C0006", name="宋体")
            elif a.severity == "中度舆情":
                sev_cell.fill = PatternFill("solid", fgColor="FFE699")
            elif a.severity == "重大利好":
                sev_cell.fill = PatternFill("solid", fgColor="C6EFCE")
                sev_cell.font = Font(bold=True, color="006100", name="宋体")
            elif a.severity == "中度利好":
                sev_cell.fill = PatternFill("solid", fgColor="E2EFDA")
            ws.cell(row=r, column=10, value=a.sentiment)
            ws.cell(row=r, column=11, value="、".join(a.keywords_matched[:6]))

        # v4.3.2: 批量设置数据区字体
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=11):
            for cell in row:
                if cell.font and cell.font.name and cell.font.name != "宋体":
                    continue  # 保留已自定义的字体（超鏈接、严重色等）
                if cell.value is not None and (cell.font is None or not cell.font.name):
                    cell.font = Font(name="宋体", size=11)

        # 设置列宽
        widths = [6, 50, 60, 16, 14, 18, 40, 24, 14, 10, 30]
        for i, w in enumerate(widths, 1):
            col_letter = chr(64 + i)
            ws.column_dimensions[col_letter].width = w

    wb.save(output_path)
    return str(output_path)


# ================================================================
# 3. Word (.docx) — 含封面 / 目录 / 章节
# ================================================================

def to_word(snapshot: SentimentSnapshot,
            output_path: Optional[Path] = None) -> str:
    """导出为 .docx，含封面 / 概览 / 分类正文。"""
    if not HAS_DOCX:
        return _placeholder("Word 导出需要 python-docx: pip install python-docx")
    if snapshot is None:
        return _placeholder("快照为空")

    output_path = Path(output_path or EXPORT_DIR / f"{snapshot.snapshot_id}.docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # v4.3.2 修复：设置默认字体为 宋体（中文）+ Times New Roman（拉丁），
    # 确保跨平台中文不乱码。
    from docx.shared import Pt
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    rFonts = rPr.makeelement(f"{{{NS}}}rFonts", {})
    rFonts.set(f"{{{NS}}}ascii", "Times New Roman")
    rFonts.set(f"{{{NS}}}hAnsi", "Times New Roman")
    rFonts.set(f"{{{NS}}}eastAsia", "宋体")
    rFonts.set(f"{{{NS}}}cs", "Times New Roman")
    rPr.insert(0, rFonts)

    # ---- 标题 ----
    title = doc.add_heading("全网舆情快照", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 元信息
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"快照ID：{snapshot.snapshot_id}    抓取时间：{snapshot.created_at}").italic = True

    # ---- 概览 ----
    doc.add_heading("1. 抓取概览", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light List"
    overview_rows = [
        ("目标过滤", json.dumps(snapshot.target_filter, ensure_ascii=False)),
        ("媒体过滤", " / ".join(snapshot.source_filter)),
        ("总条数", str(snapshot.stats.get("total", 0))),
        ("正面", str(snapshot.positive_count())),
        ("舆情（负面）", str(snapshot.negative_count())),
        ("中性", str(snapshot.neutral_count())),
        ("风险分布", json.dumps(snapshot.stats.get("by_severity", {}), ensure_ascii=False)),
    ]
    for k, v in overview_rows:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = v

    # ---- 章节 ----
    sections = [
        ("2. ✅ 正面新闻", "positive", "🟢"),
        ("3. ⚠️ 舆情（按严重度倒序）", "negative", "🔴"),
        ("4. ℹ️ 中性公告", "neutral", "⚪"),
    ]
    for sec_title, sent_type, icon in sections:
        doc.add_heading(sec_title, level=1)
        items = [a for a in snapshot.articles if a.sentiment == sent_type]
        if sent_type == "negative":
            items.sort(key=lambda a: (-a.sentiment_score, a.publish_time or ""))
        else:
            items.sort(key=lambda a: (a.publish_time or ""), reverse=True)

        if not items:
            doc.add_paragraph("（本分类暂无内容）").italic = True
            continue

        for idx, a in enumerate(items, 1):
            heading_text = f"{icon} {idx}. [{a.target_name}] 《{a.title}》"
            doc.add_heading(heading_text, level=2)

            t = doc.add_table(rows=0, cols=2)
            t.style = "Light Grid"
            for k, v in [
                ("标题", a.title),
                ("内容简介", a.summary or "(无摘要)"),
                ("发布平台", f"{a.source} ({a.source_type})"),
                ("发布时间", a.publish_time),
                ("页面连接", a.url),
                ("严重等级", a.severity),
                ("命中关键词", "、".join(a.keywords_matched[:8])),
                ("情感分数", f"{a.sentiment_score:.1f}"),
            ]:
                row = t.add_row().cells
                row[0].text = k
                row[1].text = v

    doc.save(output_path)
    return str(output_path)


# ================================================================
# 4. JSON / CSV — 备份
# ================================================================

def to_json(snapshot: SentimentSnapshot, output_path: Optional[Path] = None) -> str:
    p = Path(output_path or EXPORT_DIR / f"{snapshot.snapshot_id}.json")
    p.write_text(json.dumps(snapshot.to_dict(), ensure_ascii=False, indent=2), encoding="utf-8")
    return str(p)


def to_csv(snapshot: SentimentSnapshot, output_path: Optional[Path] = None) -> str:
    p = Path(output_path or EXPORT_DIR / f"{snapshot.snapshot_id}.csv")
    with open(p, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["标题", "内容简介", "发布平台", "媒体类型", "发布时间", "页面连接", "目标", "情感", "严重等级", "命中关键词"])
        for a in snapshot.articles:
            writer.writerow([
                a.title, a.summary, a.source, a.source_type, a.publish_time,
                a.url, f"{a.target_name} ({a.target_type})", a.sentiment, a.severity,
                "、".join(a.keywords_matched[:6])
            ])
    return str(p)


# ================================================================
# 5. 总入口 — auto 模式
# ================================================================

def export(snapshot: SentimentSnapshot,
           fmt: str = "auto",
           output_path: Optional[Path] = None) -> Dict[str, str]:
    """统一导出。
    fmt: word / excel / csv / json / dialog / all / auto
    """
    out: Dict[str, str] = {}
    if snapshot is None:
        return out
    if fmt == "dialog":
        out["dialog"] = to_dialog(snapshot)
        return out
    if fmt == "word":
        out["word"] = to_word(snapshot, output_path)
    elif fmt == "excel":
        out["excel"] = to_excel(snapshot, output_path)
    elif fmt == "csv":
        out["csv"] = to_csv(snapshot, output_path)
    elif fmt == "json":
        out["json"] = to_json(snapshot, output_path)
    elif fmt in ("all", "auto"):
        out["word"] = to_word(snapshot)
        out["excel"] = to_excel(snapshot)
        out["json"] = to_json(snapshot)
        out["csv"] = to_csv(snapshot)
        out["dialog"] = to_dialog(snapshot)
    return out


def _placeholder(msg: str) -> str:
    return f"⚠️ {msg}"


# CLI
def _cli():
    import argparse
    parser = argparse.ArgumentParser(description="cn-financial-scraper 舆情快照导出 v4.3")
    parser.add_argument("snapshot_id", type=str, help="快照ID（或传 json 文件路径）")
    parser.add_argument("--fmt", type=str, default="auto",
                        help="输出格式 word/excel/csv/json/dialog/all/auto")
    args = parser.parse_args()

    snapshot_file = SNAPSHOT_DIR if not Path(args.snapshot_id).exists() else Path(args.snapshot_id)
    # 简单解析快照
    if Path(args.snapshot_id).exists():
        raw = json.loads(Path(args.snapshot_id).read_text(encoding="utf-8"))
        from sentiment_crawler import SentimentSnapshot, SentimentArticle  # type: ignore
        arts = [SentimentArticle(**a) for a in raw.get("articles", [])]
        snap = SentimentSnapshot(
            snapshot_id=raw["snapshot_id"], created_at=raw["created_at"],
            target_filter=raw.get("target_filter", {}),
            source_filter=raw.get("source_filter", []),
            articles=arts,
            stats=raw.get("stats", {}),
        )
        snap.extra_path = args.snapshot_id  # type: ignore[attr-defined]
    else:
        print(f"⚠️ 快照文件不存在: {args.snapshot_id}")
        return

    results = export(snap, fmt=args.fmt)
    for k, v in results.items():
        if k == "dialog":
            print(v)
        else:
            print(f"✅ [{k.upper()}] {v}")


if __name__ == "__main__":
    _cli()
