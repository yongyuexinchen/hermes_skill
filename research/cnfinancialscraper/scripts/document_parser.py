# -*- coding: utf-8 -*-
"""
文档解析模块
支持PDF、Word、Excel等金融文档解析
"""

import json
import re
import os
from pathlib import Path
from typing import Dict, Any, List, Optional

# 文档解析依赖检测
PDF_AVAILABLE = False
DOCX_AVAILABLE = False
XLSX_AVAILABLE = False

try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    pass

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    pass

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    pass

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    pass


SKILL_DATA_DIR = Path(__file__).parent.parent / "data"


def parse_pdf(file_path: str) -> Dict[str, Any]:
    """
    解析PDF文档

    Args:
        file_path: PDF文件路径

    Returns:
    解析结果
    """
    result = {
        "file_type": "pdf",
        "file_path": file_path,
        "pages": 0,
        "text_content": "",
        "tables": [],
        "financial_data": {}
    }

    if not PDF_AVAILABLE:
        result["error"] = "PyPDF2未安装，请运行: pip install PyPDF2"
        return result

    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            result["pages"] = len(reader.pages)

            # 提取文本
            full_text = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    full_text.append(text)

            result["text_content"] = "\n\n".join(full_text)

            # 尝试提取财务数据
            result["financial_data"] = extract_financial_numbers(result["text_content"])

    except Exception as e:
        result["error"] = f"PDF解析失败: {e}"

    return result


def parse_docx(file_path: str) -> Dict[str, Any]:
    """
    解析Word文档

    Args:
        file_path: DOCX文件路径

    Returns:
    解析结果
    """
    result = {
        "file_type": "docx",
        "file_path": file_path,
        "paragraphs": [],
        "tables": [],
        "text_content": "",
        "financial_data": {}
    }

    if not DOCX_AVAILABLE:
        result["error"] = "python-docx未安装，请运行: pip install python-docx"
        return result

    try:
        doc = DocxDocument(file_path)

        # 提取段落
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        result["paragraphs"] = paragraphs
        result["text_content"] = "\n\n".join(paragraphs)

        # 提取表格
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        result["tables"] = tables

        # 财务数据
        result["financial_data"] = extract_financial_numbers(result["text_content"])

    except Exception as e:
        result["error"] = f"DOCX解析失败: {e}"

    return result


def parse_xlsx(file_path: str, sheet_name: Optional[str] = None) -> Dict[str, Any]:
    """
    解析Excel文档

    Args:
        file_path: Excel文件路径
        sheet_name: 指定工作表名称（默认第一个）

    Returns:
    解析结果
    """
    result = {
        "file_type": "xlsx",
        "file_path": file_path,
        "sheets": {},
        "financial_data": {}
    }

    if not XLSX_AVAILABLE:
        result["error"] = "openpyxl未安装，请运行: pip install openpyxl"
        return result

    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)

        result["sheet_names"] = wb.sheetnames

        # 读取指定sheet或全部
        sheets_to_read = [sheet_name] if sheet_name else wb.sheetnames

        for sheet in sheets_to_read:
            ws = wb[sheet]
            data = []

            for row in ws.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else "" for cell in row]
                if any(row_data):
                    data.append(row_data)

            result["sheets"][sheet] = {
                "rows": len(data),
                "data": data[:100]  # 限制前100行
            }

        # 财务数据
        all_text = json.dumps(result["sheets"])
        result["financial_data"] = extract_financial_numbers(all_text)

    except Exception as e:
        result["error"] = f"Excel解析失败: {e}"

    return result


def parse_txt(file_path: str) -> Dict[str, Any]:
    """
    解析文本文件

    Args:
        file_path: 文本文件路径

    Returns:
    解析结果
    """
    result = {
        "file_type": "txt",
        "file_path": file_path,
        "text_content": "",
        "financial_data": {}
    }

    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            result["text_content"] = f.read()

        result["financial_data"] = extract_financial_numbers(result["text_content"])

    except Exception as e:
        try:
            # 尝试GBK编码
            with open(file_path, 'r', encoding='gbk') as f:
                result["text_content"] = f.read()
            result["financial_data"] = extract_financial_numbers(result["text_content"])
        except:
            result["error"] = f"文本解析失败: {e}"

    return result


def extract_financial_numbers(text: str) -> Dict[str, Any]:
    """
    从文本中提取财务数字

    Args:
        text: 文本内容

    Returns:
    财务数据字典
    """
    data = {}

    # 提取金额（亿/万）
    amount_patterns = [
        (r'([+-]?\d+\.?\d*)\s*亿元', 'amount_yi'),
        (r'([+-]?\d+\.?\d*)\s*万元', 'amount_wan'),
        (r'([+-]?\d+\.?\d*)\s*%，', 'percentage'),
    ]

    amounts = []
    percentages = []

    for pattern, ptype in amount_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            try:
                value = float(match)
                if ptype == 'amount_yi':
                    amounts.append(value)
                elif ptype == 'percentage':
                    percentages.append(value)
            except:
                pass

    if amounts:
        data["amounts_yi"] = amounts[:10]  # 取前10个
        data["total_yi"] = sum(amounts)

    if percentages:
        data["percentages"] = percentages[:20]

    # 提取常见财务指标
    indicators = {
        '净利润': r'净利润[：:]\s*([+-]?\d+\.?\d*)',
        '营业收入': r'营业收入[：:]\s*([+-]?\d+\.?\d*)',
        '总资产': r'总资产[：:]\s*([+-]?\d+\.?\d*)',
        '净资产': r'净资产[：:]\s*([+-]?\d+\.?\d*)',
        '每股收益': r'每股收益[（(]EPS[）)][：:]\s*([+-]?\d+\.?\d*)',
        '市盈率': r'市盈率[（(]PE[）)][：:]\s*([+-]?\d+\.?\d*)',
        '市净率': r'市净率[（(]PB[）)][：:]\s*([+-]?\d+\.?\d*)',
    }

    for name, pattern in indicators.items():
        match = re.search(pattern, text)
        if match:
            try:
                data[name] = float(match.group(1))
            except:
                pass

    return data


def parse_document(file_path: str, sheet_name: Optional[str] = None) -> Dict[str, Any]:
    """
    通用文档解析接口

    Args:
        file_path: 文件路径
        sheet_name: Excel指定工作表

    Returns:
    解析结果
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    parsers = {
        '.pdf': parse_pdf,
        '.docx': parse_docx,
        '.doc': parse_docx,  # 尝试作为docx
        '.xlsx': parse_xlsx,
        '.xls': parse_xlsx,
        '.txt': parse_txt,
        '.md': parse_txt,
        '.csv': parse_txt,
    }

    parser = parsers.get(suffix)
    if not parser:
        return {"error": f"不支持的文件类型: {suffix}"}

    if suffix == '.xlsx' or suffix == '.xls':
        return parser(file_path, sheet_name)
    else:
        return parser(file_path)


def extract_key_info(parsed_doc: Dict[str, Any]) -> str:
    """
    从解析结果中提取关键信息摘要

    Args:
        parsed_doc: parse_document的返回结果

    Returns:
    关键信息摘要文本
    """
    if "error" in parsed_doc:
        return f"解析错误: {parsed_doc['error']}"

    summary = f"【文档解析结果】\n\n"
    summary += f"文件类型: {parsed_doc.get('file_type', 'unknown').upper()}\n"

    if "pages" in parsed_doc:
        summary += f"页数: {parsed_doc['pages']}\n"
    if "paragraphs" in parsed_doc:
        summary += f"段落数: {len(parsed_doc['paragraphs'])}\n"
    if "tables" in parsed_doc:
        summary += f"表格数: {len(parsed_doc['tables'])}\n"
    if "sheets" in parsed_doc:
        summary += f"工作表: {', '.join(parsed_doc.get('sheet_names', []))}\n"

    # 财务数据摘要
    fin_data = parsed_doc.get("financial_data", {})
    if fin_data:
        summary += "\n【提取的财务数据】\n"
        for key, value in fin_data.items():
            if key == "percentages":
                summary += f"百分比数据: {value[:5]}...\n" if len(value) > 5 else f"百分比数据: {value}\n"
            elif key == "amounts_yi":
                summary += f"金额数据(亿元): {value[:5]}...\n" if len(value) > 5 else f"金额数据(亿元): {value}\n"
            elif key == "total_yi":
                summary += f"总金额: {value:.2f}亿元\n"
            else:
                summary += f"{key}: {value}\n"

    # 文本预览
    text = parsed_doc.get("text_content", "")
    if text:
        preview = text[:500].replace('\n', ' ')
        summary += f"\n【内容预览】\n{preview}..."

    return summary


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("用法: python document_parser.py <文件路径>")
        print("示例: python document_parser.py ./report.pdf")
        print("\n支持格式: PDF, DOCX, XLSX, TXT, MD, CSV")
        sys.exit(1)

    file_path = sys.argv[1]

    if not os.path.exists(file_path):
        print(f"文件不存在: {file_path}")
        sys.exit(1)

    print(f"正在解析: {file_path}\n")
    result = parse_document(file_path)
    print(extract_key_info(result))